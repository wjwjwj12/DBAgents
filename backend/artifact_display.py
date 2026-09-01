import os
import re
import shutil
import subprocess
import tempfile
import logging
import uuid
from pathlib import Path
from urllib.parse import urlsplit
from xmlrpc.client import Binary, SafeTransport, ServerProxy, Transport

import httpx


PPTX_MIME = "application/vnd.openxmlformats-officedocument.presentationml.presentation"
_ARTIFACT_SOURCE_MARKERS = ("[历史产物:", "<!DOCTYPE", "<html", "<style")
_PROCESS_PREAMBLE_PREFIXES = ("我来", "让我", "我先", "我会先", "接下来我", "下面我")
logger = logging.getLogger(__name__)


def artifact_preview_path(storage_path: str) -> str:
    return f"{storage_path}.preview.md"


def artifact_pdf_preview_path(storage_path: str) -> str:
    return f"{storage_path}.preview.pdf"


def libreoffice_executable() -> str | None:
    return shutil.which("libreoffice") or shutil.which("soffice")


def libreoffice_convert_url() -> str:
    return os.getenv("LIBREOFFICE_CONVERT_URL", "").strip().rstrip("/")


def unoserver_config() -> tuple[str, int, str] | None:
    host = os.getenv("LIBREOFFICE_UNOSERVER_HOST", "").strip()
    if not host:
        return None
    port = int(os.getenv("LIBREOFFICE_UNOSERVER_PORT", "2003"))
    protocol = os.getenv("LIBREOFFICE_UNOSERVER_PROTOCOL", "http").strip().lower()
    if not 1 <= port <= 65535 or protocol not in {"http", "https"}:
        raise ValueError("Invalid LibreOffice unoserver connection configuration")
    return host, port, protocol


def pptx_pdf_preview_available(storage_path: str) -> bool:
    return bool(
        os.path.exists(artifact_pdf_preview_path(storage_path))
        or unoserver_config()
        or libreoffice_convert_url()
        or libreoffice_executable() is not None
    )


class _TimeoutTransport(Transport):
    def __init__(self, timeout: int):
        super().__init__()
        self.timeout = timeout

    def make_connection(self, host):
        connection = super().make_connection(host)
        connection.timeout = self.timeout
        return connection


class _TimeoutSafeTransport(SafeTransport):
    def __init__(self, timeout: int):
        super().__init__()
        self.timeout = timeout

    def make_connection(self, host):
        connection = super().make_connection(host)
        connection.timeout = self.timeout
        return connection


def _convert_pptx_via_unoserver(source: Path, target: Path, timeout: int) -> str:
    config = unoserver_config()
    if config is None:
        raise RuntimeError("LibreOffice unoserver is not configured")
    host, port, protocol = config
    temporary = target.with_name(f".{target.stem}.{uuid.uuid4().hex}.pdf")
    try:
        transport = _TimeoutSafeTransport(timeout) if protocol == "https" else _TimeoutTransport(timeout)
        with ServerProxy(
            f"{protocol}://{host}:{port}",
            allow_none=True,
            transport=transport,
        ) as proxy:
            proxy.info()
            result = proxy.convert(
                None,
                Binary(source.read_bytes()),
                None,
                "pdf",
                None,
                [],
                True,
                None,
                None,
            )
        converted = result.data if isinstance(result, Binary) else bytes(result)
        if not converted.startswith(b"%PDF-"):
            raise RuntimeError("unoserver did not return a PDF")
        temporary.write_bytes(converted)
        os.replace(temporary, target)
        return str(target)
    except RuntimeError:
        raise
    except Exception as exc:
        raise RuntimeError(f"LibreOffice unoserver conversion failed: {exc}") from exc
    finally:
        temporary.unlink(missing_ok=True)


def _remote_convert_candidates(configured_url: str) -> list[tuple[str, str]]:
    path = urlsplit(configured_url).path.rstrip("/")
    if path:
        return [(configured_url, "files"), (configured_url, "file")]
    return [
        (f"{configured_url}/forms/libreoffice/convert", "files"),
        (f"{configured_url}/convert", "file"),
        (f"{configured_url}/api/convert", "file"),
    ]


def _convert_pptx_remotely(source: Path, target: Path, configured_url: str, timeout: int) -> str:
    errors = []
    for url, field_name in _remote_convert_candidates(configured_url):
        try:
            with source.open("rb") as stream:
                response = httpx.post(
                    url,
                    files={field_name: (source.name, stream, PPTX_MIME)},
                    data={"format": "pdf"},
                    timeout=httpx.Timeout(timeout, connect=min(5, timeout)),
                )
            if response.status_code == 404:
                errors.append(f"{url}: HTTP 404")
                continue
            response.raise_for_status()
            if not response.content.startswith(b"%PDF-"):
                raise RuntimeError(f"{url} did not return a PDF")
            temporary = target.with_name(f".{target.name}.{uuid.uuid4().hex}.tmp")
            try:
                temporary.write_bytes(response.content)
                os.replace(temporary, target)
            finally:
                temporary.unlink(missing_ok=True)
            return str(target)
        except (OSError, httpx.HTTPError, RuntimeError) as exc:
            errors.append(f"{url}: {exc}")
    raise RuntimeError("Remote LibreOffice conversion failed: " + "; ".join(errors))


def ensure_pptx_pdf_preview(storage_path: str) -> str | None:
    source = Path(storage_path)
    target = Path(artifact_pdf_preview_path(storage_path))
    if target.is_file() and target.stat().st_mtime_ns >= source.stat().st_mtime_ns:
        return str(target)
    timeout = max(10, int(os.getenv("LIBREOFFICE_CONVERT_TIMEOUT_SECONDS", "120")))
    uno_error = None
    if unoserver_config() is not None:
        try:
            return _convert_pptx_via_unoserver(source, target, timeout)
        except RuntimeError as exc:
            uno_error = exc
            logger.warning("Remote unoserver preview conversion failed; trying other converters: %s", exc)
    remote_url = libreoffice_convert_url()
    remote_error = None
    if remote_url:
        try:
            return _convert_pptx_remotely(source, target, remote_url, timeout)
        except RuntimeError as exc:
            remote_error = exc
            logger.warning("Remote LibreOffice preview conversion failed; trying local executable: %s", exc)
    executable = libreoffice_executable()
    if executable is None:
        if remote_error is not None:
            raise remote_error
        if uno_error is not None:
            raise uno_error
        return None
    with tempfile.TemporaryDirectory(prefix="pptx-preview-", dir=str(source.parent)) as output_dir:
        profile_dir = Path(output_dir) / "profile"
        profile_dir.mkdir()
        completed = subprocess.run(
            [
                executable,
                f"-env:UserInstallation={profile_dir.as_uri()}",
                "--headless",
                "--convert-to",
                "pdf",
                "--outdir",
                output_dir,
                str(source),
            ],
            check=False,
            capture_output=True,
            text=True,
            timeout=timeout,
        )
        generated = Path(output_dir) / f"{source.stem}.pdf"
        if completed.returncode != 0 or not generated.is_file():
            details = (completed.stderr or completed.stdout or "unknown conversion error").strip()
            raise RuntimeError(f"LibreOffice failed to convert PPTX preview: {details}")
        os.replace(generated, target)
    return str(target)


def read_artifact_preview(artifact) -> tuple[str, str]:
    if artifact.mime_type == "text/html":
        path = artifact.storage_path
        preview_kind = "html"
    elif artifact.mime_type == "application/pdf":
        return "pdf", ""
    elif artifact.mime_type.startswith("image/"):
        return "image", ""
    elif artifact.mime_type == PPTX_MIME:
        if pptx_pdf_preview_available(artifact.storage_path):
            return "pdf", ""
        path = artifact_preview_path(artifact.storage_path)
        preview_kind = "markdown"
    elif artifact.mime_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        path = artifact_preview_path(artifact.storage_path)
        preview_kind = "markdown"
        if not os.path.exists(path) and os.path.exists(artifact.storage_path):
            from docx import Document

            document = Document(artifact.storage_path)
            blocks = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
            for table in document.tables:
                for row in table.rows:
                    blocks.append("| " + " | ".join(cell.text.replace("\n", " ").strip() for cell in row.cells) + " |")
            return "text", "\n".join(blocks)
    else:
        path = artifact.storage_path
        preview_kind = "text"
    if not os.path.exists(path):
        return preview_kind, ""
    with open(path, "r", encoding="utf-8") as source:
        return preview_kind, source.read()


def sanitize_artifact_final_text(text: str) -> str:
    indexes = [text.find(marker) for marker in _ARTIFACT_SOURCE_MARKERS if text.find(marker) >= 0]
    return text[:min(indexes)].strip() if indexes else text.strip()


def contains_artifact_source(text: str) -> bool:
    return any(marker in text for marker in _ARTIFACT_SOURCE_MARKERS)


def is_process_preamble(text: str) -> bool:
    normalized = text.strip()
    return any(normalized.startswith(prefix) for prefix in _PROCESS_PREAMBLE_PREFIXES)


def could_be_process_preamble(text: str) -> bool:
    normalized = text.strip()
    return bool(normalized) and any(
        prefix.startswith(normalized) or normalized.startswith(prefix)
        for prefix in _PROCESS_PREAMBLE_PREFIXES
    )


def split_leading_process_preamble(text: str) -> tuple[list[str], str]:
    paragraphs = [part.strip() for part in re.split(r"\n\s*\n", text.strip()) if part.strip()]
    preamble = []
    while paragraphs and is_process_preamble(paragraphs[0]):
        preamble.append(paragraphs.pop(0))
    return preamble, "\n\n".join(paragraphs)
