import os
import re

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


BODY_FONT = "宋体"
HEADING_FONT = "微软雅黑"
ACCENT = RGBColor(31, 78, 120)


def _set_run_font(run, font_name: str, size: Pt | None = None, color=None, bold=None, italic=None):
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    run._element.rPr.rFonts.set(qn("w:ascii"), font_name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), font_name)
    if size is not None:
        run.font.size = size
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.font.bold = bold
    if italic is not None:
        run.font.italic = italic


def _set_style_font(style, font_name: str, size: Pt, color=None, bold=None):
    style.font.name = font_name
    style._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    style._element.rPr.rFonts.set(qn("w:ascii"), font_name)
    style._element.rPr.rFonts.set(qn("w:hAnsi"), font_name)
    style.font.size = size
    if color is not None:
        style.font.color.rgb = color
    if bold is not None:
        style.font.bold = bold


def _configure_document(doc: Document) -> None:
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        section.header_distance = Inches(0.49)
        section.footer_distance = Inches(0.49)
        header = section.header.paragraphs[0]
        header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        header.paragraph_format.space_after = Pt(0)
        _set_run_font(header.add_run("auto-agent 生成文档"), BODY_FONT, Pt(9), RGBColor(128, 128, 128))

    normal = doc.styles["Normal"]
    _set_style_font(normal, BODY_FONT, Pt(11))
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for style_name, size, color in (
        ("Heading 1", Pt(16), ACCENT),
        ("Heading 2", Pt(13), ACCENT),
        ("Heading 3", Pt(12), RGBColor(31, 58, 95)),
    ):
        style = doc.styles[style_name]
        _set_style_font(style, HEADING_FONT, size, color, True)
        style.paragraph_format.space_before = Pt(14 if style_name == "Heading 1" else 10)
        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.keep_with_next = True

    for style_name in ("List Bullet", "List Number"):
        style = doc.styles[style_name]
        _set_style_font(style, BODY_FONT, Pt(11))
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25


def _set_cell_background(cell, fill_hex: str) -> None:
    tc_pr = cell._element.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill_hex)
    tc_pr.append(shd)


def _write_inline(paragraph, text: str, *, font_name: str = BODY_FONT, size: Pt = Pt(11), color=None) -> None:
    """Write a safe subset of Markdown inline syntax without dropping unmatched text."""
    token_pattern = re.compile(r"(\*\*.+?\*\*|__.+?__|`.+?`|\[.+?\]\(.+?\))")
    position = 0
    for match in token_pattern.finditer(text):
        if match.start() > position:
            _set_run_font(paragraph.add_run(text[position:match.start()]), font_name, size, color)
        token = match.group(0)
        if token.startswith(("**", "__")):
            run = paragraph.add_run(token[2:-2])
            _set_run_font(run, font_name, size, color, bold=True)
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            _set_run_font(run, "Consolas", size, color)
            run.font.highlight_color = 16
        else:
            label, url = re.match(r"\[(.*?)\]\((.*?)\)", token).groups()
            run = paragraph.add_run(f"{label} ({url})")
            _set_run_font(run, font_name, size, color)
            run.font.underline = True
        position = match.end()
    if position < len(text) or not text:
        _set_run_font(paragraph.add_run(text[position:]), font_name, size, color)


def _is_table_separator(line: str) -> bool:
    cells = [cell.strip() for cell in line.strip().strip("|").split("|")]
    return bool(cells) and all(re.fullmatch(r":?-{3,}:?", cell) for cell in cells)


def _flush_table(doc: Document, lines: list[str]) -> None:
    if len(lines) < 2 or not _is_table_separator(lines[1]):
        for line in lines:
            paragraph = doc.add_paragraph()
            _write_inline(paragraph, line)
        return
    rows = [
        [cell.strip() for cell in line.strip().strip("|").split("|")]
        for index, line in enumerate(lines)
        if index != 1
    ]
    columns = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=columns)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for row_index, row in enumerate(rows):
        for column_index in range(columns):
            cell = table.cell(row_index, column_index)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(2)
            _write_inline(paragraph, row[column_index] if column_index < len(row) else "", size=Pt(10.5))
            if row_index == 0:
                _set_cell_background(cell, "E8EEF5")
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.color.rgb = ACCENT


def export_markdown_to_docx(markdown_content: str, output_path: str, title: str = "文档产物"):
    """Export Markdown to a Chinese Word document while preserving all source text."""
    doc = Document()
    _configure_document(doc)

    title_paragraph = doc.add_paragraph()
    title_paragraph.paragraph_format.space_after = Pt(14)
    _set_run_font(title_paragraph.add_run(title), HEADING_FONT, Pt(20), ACCENT, bold=True)

    lines = markdown_content.replace("\r\n", "\n").replace("\r", "\n").split("\n")
    table_lines: list[str] = []

    def flush_table():
        nonlocal table_lines
        if table_lines:
            _flush_table(doc, table_lines)
            table_lines = []

    for line in lines:
        stripped = line.strip()
        if stripped.startswith("|") and stripped.endswith("|"):
            table_lines.append(stripped)
            continue
        flush_table()

        if not stripped:
            doc.add_paragraph()
            continue
        heading = re.match(r"^(#{1,3})\s+(.+)$", stripped)
        if heading:
            level = len(heading.group(1))
            paragraph = doc.add_paragraph(style=f"Heading {level}")
            _write_inline(paragraph, heading.group(2), font_name=HEADING_FONT, size=(Pt(16), Pt(13), Pt(12))[level - 1], color=(ACCENT if level < 3 else RGBColor(31, 58, 95)))
            continue
        if re.fullmatch(r"(?:-{3,}|\*{3,}|_{3,})", stripped):
            paragraph = doc.add_paragraph()
            paragraph.paragraph_format.space_before = Pt(4)
            paragraph.paragraph_format.space_after = Pt(4)
            _set_run_font(paragraph.add_run("─" * 42), BODY_FONT, Pt(8), RGBColor(180, 190, 200))
            continue
        quote = re.match(r"^>\s?(.*)$", stripped)
        if quote:
            paragraph = doc.add_paragraph()
            paragraph.paragraph_format.left_indent = Inches(0.32)
            paragraph.paragraph_format.space_after = Pt(6)
            _write_inline(paragraph, quote.group(1), font_name=BODY_FONT, size=Pt(11), color=RGBColor(90, 100, 110))
            continue
        bullet = re.match(r"^[-*+]\s+(.+)$", stripped)
        if bullet:
            paragraph = doc.add_paragraph(style="List Bullet")
            _write_inline(paragraph, bullet.group(1))
            continue
        numbered = re.match(r"^\d+[.)]\s+(.+)$", stripped)
        if numbered:
            paragraph = doc.add_paragraph(style="List Number")
            _write_inline(paragraph, numbered.group(1))
            continue
        paragraph = doc.add_paragraph()
        _write_inline(paragraph, stripped)

    flush_table()
    output_dir = os.path.dirname(output_path)
    if output_dir:
        os.makedirs(output_dir, exist_ok=True)
    doc.save(output_path)
    return output_path
