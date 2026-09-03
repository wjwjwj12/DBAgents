import os
import shutil
import asyncio
import io
import json
import logging
import re
import subprocess
import time
import uuid
import httpx
from dotenv import load_dotenv
if os.getenv("APP_ENV", "development").lower() != "production":
    load_dotenv()

from fastapi import Depends, FastAPI, UploadFile, File, Form, HTTPException, Request
from fastapi.responses import StreamingResponse, JSONResponse, FileResponse, RedirectResponse
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel, Field
from sqlalchemy import func, select
from typing import Optional

from database import init_db, SessionLocal
from harness.thread_state import ThreadRecorder, get_or_create_state, update_state
from models import ArtifactModel, AttachmentModel, ConversationModel, ConversationStateModel, MessageModel, PlanTaskModel, RunEventModel, RunModel, ThreadEventModel, UserModel, utc_now
from security import create_download_url, create_preview_url, verify_download_token
from runtime_paths import ATTACHMENT_DIR as RUNTIME_ATTACHMENT_DIR
from capabilities.skill_installer import (
    SkillPackageError,
    install_preflighted_skill,
    preflight_skill_archive,
)
from capabilities.skill_registry import SkillRegistry, get_user_skill_root
from logging_config import configure_logging
from auth import (
    RequestIdentity,
    get_or_create_user,
    get_request_identity,
    normalize_identity,
    owned_artifact,
    owned_attachment,
    owned_conversation,
    owned_run,
    validate_auth_configuration,
)
from artifact_display import (
    PPTX_MIME,
    contains_artifact_source,
    ensure_pptx_pdf_preview,
    read_artifact_preview,
    sanitize_artifact_final_text,
    split_leading_process_preamble,
)


APP_LOG_FILE = configure_logging()
logger = logging.getLogger(__name__)


def _groups_from_run_events(run_id: str, events) -> list[dict]:
    groups = []
    pending_thought_parts = []
    pending_answer_parts = []
    pending_thought_group = None

    def add_group(title, thought=None, created_at=None):
        group = {
            "id": f"{run_id}-{len(groups) + 1}",
            "title": title,
            "thoughts": [],
            "actions": [],
            "started_at": created_at.isoformat() if created_at else None,
        }
        if thought:
            group["thoughts"].append(thought)
        groups.append(group)
        return group

    def current_group():
        return groups[-1] if groups else add_group("任务执行")

    def flush_pending_model_as_thought():
        nonlocal pending_thought_group
        reasoning_text = "".join(pending_thought_parts).strip()
        answer_text = "".join(pending_answer_parts).strip()
        pending_thought_parts.clear()
        pending_answer_parts.clear()
        if pending_thought_group is not None:
            if reasoning_text:
                pending_thought_group["thoughts"].append(reasoning_text)
            if answer_text:
                pending_thought_group["thoughts"].append(answer_text)
        pending_thought_group = None

    def add_tool_action(group, text, name):
        group["actions"].append({"text": text, "status": "loading", "_tool_name": name})

    def complete_tool_action(group, name):
        for action in reversed(group["actions"]):
            if action.get("_tool_name") == name and action["status"] == "loading":
                action["status"] = "done"
                return
        group["actions"].append({"text": f"{name} 执行完成", "status": "done"})

    for event in events:
        payload = event.payload or {}
        if event.event_type == "plan_created":
            group = add_group(payload.get("title") or "任务计划", created_at=event.created_at)
            for step in payload.get("steps") or []:
                title = step.get("title") if isinstance(step, dict) else str(step)
                group["actions"].append({"text": title, "status": "loading"})
        elif event.event_type == "model_started":
            flush_pending_model_as_thought()
            turn = int(payload.get("turn", 1))
            pending_thought_group = add_group(
                f"第 {turn} 轮思考",
                created_at=event.created_at,
            )
        elif event.event_type == "reasoning_delta":
            pending_thought_parts.append(str(payload.get("text") or ""))
        elif event.event_type == "model_delta":
            pending_answer_parts.append(str(payload.get("text") or ""))
        elif event.event_type == "tool_started":
            flush_pending_model_as_thought()
            name = payload.get("name") or "工具"
            arguments = payload.get("arguments") or {}
            if name == "update_plan":
                group = add_group(arguments.get("title") or "动态任务计划")
                for step in arguments.get("steps") or []:
                    title = step.get("title") if isinstance(step, dict) else str(step)
                    group["actions"].append({"text": title, "status": "loading"})
            elif name == "create_document":
                group = add_group("正在生成任务产物")
                add_tool_action(group, f"正在执行 {name}", name)
            else:
                add_tool_action(current_group(), f"正在执行 {name}", name)
        elif event.event_type == "tool_completed":
            name = payload.get("name") or "工具"
            complete_tool_action(current_group(), name)
        elif event.event_type == "model_empty":
            current_group()["actions"].append({"text": "模型未返回有效内容，正在自动恢复", "status": "loading"})
        elif event.event_type == "tool_approval_required":
            flush_pending_model_as_thought()
    reasoning_text = "".join(pending_thought_parts).strip()
    process_preamble, _answer_text = split_leading_process_preamble("".join(pending_answer_parts))
    if reasoning_text and pending_thought_group is not None:
        pending_thought_group["thoughts"].append(reasoning_text)
    if pending_thought_group is not None:
        pending_thought_group["thoughts"].extend(process_preamble)
    if pending_thought_group is not None and not pending_thought_group["thoughts"] and not pending_thought_group["actions"]:
        groups.remove(pending_thought_group)
    for group in groups:
        for action in group["actions"]:
            action.pop("_tool_name", None)
    return groups


def _restore_run_groups(db, run_id: str) -> list[dict]:
    events = (
        db.query(RunEventModel)
        .filter(RunEventModel.run_id == run_id)
        .order_by(RunEventModel.sequence)
        .all()
    )
    return _groups_from_run_events(run_id, events)


def _artifact_download_filename(title: str, storage_path: str) -> str:
    suffix = os.path.splitext(storage_path)[1]
    safe_title = re.sub(r'[<>:"/\\|?*\x00-\x1f]', "_", title).strip(" .") or "artifact"
    if suffix and not safe_title.lower().endswith(suffix.lower()):
        safe_title += suffix
    return safe_title

app = FastAPI(title="交数智航 API")


def _finalize_interrupted_runs() -> int:
    db = SessionLocal()
    try:
        runs = db.query(RunModel).filter(RunModel.status.in_({"pending", "running"})).all()
        if not runs:
            return 0
        completed_at = utc_now()
        run_ids = [run.id for run in runs]
        for run in runs:
            run.status = "failed"
            run.error_message = "服务重启或上一次执行中断，请重试该任务。"
            run.completed_at = completed_at
        for task in db.query(PlanTaskModel).filter(
            PlanTaskModel.run_id.in_(run_ids),
            PlanTaskModel.status.in_({"pending", "running"}),
        ):
            task.status = "failed"
            task.error_message = "任务执行进程已中断"
            task.completed_at = completed_at
        db.commit()
        return len(runs)
    finally:
        db.close()


# Initialize SQLite/Database schema on startup
@app.on_event("startup")
def on_startup():
    validate_auth_configuration()
    init_db()
    interrupted = _finalize_interrupted_runs()
    if interrupted:
        logger.warning("Marked interrupted runs as failed count=%d", interrupted)
    logger.info("Application started log_file=%s", APP_LOG_FILE)


@app.on_event("shutdown")
async def on_shutdown():
    from orchestration.checkpoint import close_checkpointer
    await close_checkpointer()

DEFAULT_ALLOWED_ORIGINS = "http://localhost:6477,http://127.0.0.1:6477"
ALLOWED_ORIGINS = list(dict.fromkeys(
    origin.strip()
    for origin in f"{DEFAULT_ALLOWED_ORIGINS},{os.getenv('ALLOWED_ORIGINS', '')}".split(",")
    if origin.strip()
))
MAX_UPLOAD_BYTES = int(os.getenv("MAX_UPLOAD_BYTES", str(100 * 1024 * 1024)))
MAX_SKILL_UPLOAD_BYTES = int(os.getenv("MAX_SKILL_UPLOAD_BYTES", str(20 * 1024 * 1024)))
TEXT_UPLOAD_EXTENSIONS = {"pdf", "docx", "txt", "md", "markdown"}
BINARY_UPLOAD_EXTENSIONS = {"pptx", "mp3", "m4a", "wav", "aac", "ogg", "flac"}
ALLOWED_UPLOAD_EXTENSIONS = TEXT_UPLOAD_EXTENSIONS | BINARY_UPLOAD_EXTENSIONS

app.add_middleware(
    CORSMiddleware,
    allow_origins=ALLOWED_ORIGINS,
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


@app.exception_handler(Exception)
async def unhandled_exception_handler(request: Request, exc: Exception):
    logger.exception("Unhandled request error method=%s path=%s", request.method, request.url.path)
    return JSONResponse(status_code=500, content={"detail": "服务器内部错误，请查看运行日志"})

FILE_PARSE_URL = os.getenv("FILE_PARSE_URL")
FILE_PARSE_TOKEN = os.getenv("FILE_PARSE_TOKEN")
FILE_PARSE_BACKEND = os.getenv("FILE_PARSE_BACKEND", "pipeline")
FILE_PARSE_TIMEOUT_SECONDS = float(os.getenv("FILE_PARSE_TIMEOUT_SECONDS", "600"))
FRONTEND_URL = os.getenv("FRONTEND_URL", "http://localhost:6477").rstrip("/")
ATTACHMENT_DIR = str(RUNTIME_ATTACHMENT_DIR)
os.makedirs(ATTACHMENT_DIR, exist_ok=True)


@app.get("/", include_in_schema=False)
async def frontend_entry():
    return RedirectResponse(FRONTEND_URL)

class ChatRequest(BaseModel):
    query: str = Field(min_length=1, max_length=20_000)
    context_text: str = Field(default="", max_length=2_000_000)
    conversation_id: str = Field(default="default_conv", min_length=1, max_length=36)
    attachment_name: Optional[str] = Field(default=None, max_length=200)
    attachment_id: Optional[str] = Field(default=None, max_length=36)
    attachment_ids: list[str] = Field(default_factory=list, max_length=20)
    selected_skill_ids: list[str] = Field(default_factory=list, max_length=1)


class ConversationStateRequest(BaseModel):
    is_archived: Optional[bool] = None
    is_pinned: Optional[bool] = None


def _accessible_skill_registry(identity: RequestIdentity, load_instructions: bool = True) -> SkillRegistry:
    from agent import skill_registry

    user_root = get_user_skill_root(identity.tenant_id, identity.user_id)
    return SkillRegistry(skill_registry.skill_root, [user_root], load_instructions=load_instructions)


def _local_skill_intro(candidate) -> str:
    description = str(candidate.metadata.get("description") or "").strip()
    description = re.sub(r"[`#*_]+", "", description)
    description = re.sub(r"\s+", " ", description).strip(" ，。；;：:")
    first_sentence = re.split(r"[。！？!?；;]", description, maxsplit=1)[0].strip()
    return (first_sentence or "提供可按需加载的专业任务处理能力")[:48]


async def _generate_skill_intro(candidate) -> str:
    fallback = _local_skill_intro(candidate)
    manifest_path = next((path for path in candidate.files if path.name.lower() == "manifest.json"), None)
    entrypoint = None
    if manifest_path is not None:
        try:
            manifest = json.loads(candidate.files[manifest_path].decode("utf-8"))
            requested_entrypoint = str(manifest.get("entrypoint") or "SKILL.md").replace("\\", "/")
            entrypoint = next((path for path in candidate.files if path.as_posix() == requested_entrypoint), None)
        except (UnicodeDecodeError, json.JSONDecodeError, AttributeError):
            entrypoint = None
    entrypoint = entrypoint or next((path for path in candidate.files if path.name.lower() == "skill.md"), None)
    if entrypoint is None:
        return fallback
    try:
        instructions = candidate.files[entrypoint].decode("utf-8")[:12_000]
    except UnicodeDecodeError:
        return fallback
    try:
        from agent import _get_llm_client

        response = await asyncio.wait_for(
            _get_llm_client().ainvoke([
                {
                    "role": "system",
                    "content": "将 Skill 说明提炼为中文简介。只输出一句话，18-36 个汉字，不写工具名、版本号、执行步骤、Markdown 或宣传用语。",
                },
                {"role": "user", "content": instructions},
            ]),
            timeout=15,
        )
        content = response.content
        if isinstance(content, list):
            content = "".join(str(item.get("text") or "") for item in content if isinstance(item, dict))
        summary = re.sub(r"\s+", " ", str(content or "")).strip(" `#*_，。；;：:\"'")
        return summary[:48] if 8 <= len(summary) <= 80 else fallback
    except Exception as exc:
        logger.warning("Skill intro generation failed; using local fallback: %s", exc)
        return fallback


@app.get("/api/v1/skills")
async def list_skills(identity: RequestIdentity = Depends(get_request_identity)):
    identity = normalize_identity(identity)
    return {"skills": _accessible_skill_registry(identity, load_instructions=False).list_skills()}


@app.post("/api/v1/skills/upload")
async def upload_skill(
    file: UploadFile = File(...),
    identity: RequestIdentity = Depends(get_request_identity),
):
    identity = normalize_identity(identity)
    from agent import skill_registry, tool_registry

    registry = _accessible_skill_registry(identity, load_instructions=False)
    existing = set()
    for skill in registry.list_skills():
        existing.add(str(skill["id"]))
        existing.update(str(alias) for alias in skill["aliases"])
    try:
        filename = file.filename or ""
        if not filename.lower().endswith(".zip"):
            raise HTTPException(status_code=415, detail="技能仅支持 ZIP 压缩包")
        content = await file.read(MAX_SKILL_UPLOAD_BYTES + 1)
        if len(content) > MAX_SKILL_UPLOAD_BYTES:
            raise HTTPException(status_code=413, detail="技能包超过上传大小限制")
        candidate = await asyncio.to_thread(
            preflight_skill_archive, content, tool_registry.names(), existing, filename,
        )
        candidate.metadata["description"] = await _generate_skill_intro(candidate)
        user_root = get_user_skill_root(identity.tenant_id, identity.user_id)
        installed = await asyncio.to_thread(install_preflighted_skill, candidate, user_root)
    except SkillPackageError as exc:
        raise HTTPException(status_code=422, detail=f"技能预检未通过：{exc}") from exc
    skill_registry.refresh()
    return {"success": True, "skill": installed}


class ConversationUpdateRequest(BaseModel):
    title: str = Field(min_length=1, max_length=200)


class ApprovalRequest(BaseModel):
    decision: Optional[str] = Field(default=None, pattern="^(approve|reject)$")
    response: Optional[str] = Field(default=None, min_length=1, max_length=4000)


_APPROVAL_NEGATIVE_MARKERS = (
    "不同意", "不可以", "不行", "不要", "不能", "不确认", "暂不", "先不", "先别", "拒绝", "取消", "停止", "等等",
    "修改", "调整", "改成", "但是", "不过", "有问题", "再想想",
)
_APPROVAL_AFFIRMATIVE_MARKERS = (
    "同意", "确认", "继续", "执行", "开始", "没问题", "就这样", "按这个方案",
    "按该方案", "按上述方案", "照这个方案", "好的", "好吧", "可以", "行",
)


def _approval_decision_from_response(response: str) -> str:
    normalized = re.sub(r"\s+", "", response).lower()
    if not normalized:
        raise ValueError("确认回复不能为空")
    if any(marker in normalized for marker in _APPROVAL_NEGATIVE_MARKERS):
        return "reject"
    if "?" in normalized or "？" in normalized or "是否" in normalized:
        return "reject"
    if any(marker in normalized for marker in _APPROVAL_AFFIRMATIVE_MARKERS):
        return "approve"
    return "reject"


def _parse_file_locally(file_ext: str, file_content: bytes) -> str:
    try:
        if file_ext == "pdf":
            try:
                import pymupdf
            except ImportError as exc:
                raise HTTPException(
                    status_code=503,
                    detail="PDF 解析组件未安装，请安装项目 requirements.txt 中的 PyMuPDF。",
                ) from exc
            with pymupdf.open(stream=file_content, filetype="pdf") as pdf_doc:
                return "\n".join(page.get_text() for page in pdf_doc)
        if file_ext == "docx":
            from docx import Document
            doc = Document(io.BytesIO(file_content))
            blocks = [para.text for para in doc.paragraphs if para.text.strip()]
            for table in doc.tables:
                for row in table.rows:
                    cells = [cell.text.replace("\n", " ").strip() for cell in row.cells]
                    blocks.append("| " + " | ".join(cells) + " |")
            return "\n".join(blocks)
        return file_content.decode("utf-8")
    except HTTPException:
        raise
    except (UnicodeDecodeError, ValueError, RuntimeError) as exc:
        raise HTTPException(
            status_code=422,
            detail="文件内容无法解析，可能文件已损坏、加密或格式不正确。",
        ) from exc
    except Exception as exc:
        raise HTTPException(
            status_code=422,
            detail="文件内容无法解析，可能文件已损坏、加密或格式不正确。",
        ) from exc

def _ensure_local_conversation(
    db,
    conversation_id: str,
    title: str = "新对话",
    identity: RequestIdentity | None = None,
):
    identity = normalize_identity(identity)
    user = get_or_create_user(db, identity)
    conversation = owned_conversation(db, conversation_id, identity)
    if db.get(ConversationModel, conversation_id) is not None and conversation is None:
        raise HTTPException(status_code=404, detail="Conversation not found")
    if conversation is None:
        conversation = ConversationModel(id=conversation_id, user_id=user.id, title=title)
        db.add(conversation)
        db.flush()
    get_or_create_state(db, conversation.id)
    return conversation


@app.post("/upload")
async def upload_file(
    file: UploadFile = File(...),
    conversation_id: str = Form(default="default_conv"),
    identity: RequestIdentity = Depends(get_request_identity),
):
    try:
        # Direct function calls in tests do not have FastAPI's dependency injection.
        if not isinstance(conversation_id, str):
            conversation_id = "default_conv"
        filename = file.filename or ""
        file_ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
        if file_ext not in ALLOWED_UPLOAD_EXTENSIONS:
            raise HTTPException(status_code=415, detail="Unsupported file type")

        file_content = await file.read(MAX_UPLOAD_BYTES + 1)
        if len(file_content) > MAX_UPLOAD_BYTES:
            raise HTTPException(status_code=413, detail="File exceeds upload size limit")
            
        if file_ext in BINARY_UPLOAD_EXTENSIONS:
            fallback_text = ""
        elif not FILE_PARSE_URL or not FILE_PARSE_TOKEN or "your-file-parse" in FILE_PARSE_URL:
            fallback_text = await asyncio.to_thread(_parse_file_locally, file_ext, file_content)
        else:
            files = {"files": (filename, file_content, file.content_type)}
            data = {
                "backend": FILE_PARSE_BACKEND,
                "lang_list": "ch",
                "parse_method": "auto",
                "return_md": "true",
            }
            headers = {"Authorization": f"Bearer {FILE_PARSE_TOKEN}"}
            parse_started_at = time.monotonic()
            logger.info(
                "Remote file parsing started filename=%r size_bytes=%s backend=%s timeout_seconds=%s",
                filename,
                len(file_content),
                FILE_PARSE_BACKEND,
                FILE_PARSE_TIMEOUT_SECONDS,
            )
            try:
                async with httpx.AsyncClient(timeout=FILE_PARSE_TIMEOUT_SECONDS) as client:
                    response = await client.post(FILE_PARSE_URL, files=files, data=data, headers=headers)
                    response.raise_for_status()
                    parsed = response.json()
                    fallback_text = next(
                        (str(value.get("md_content", "")) for value in parsed.get("results", {}).values()),
                        "",
                    )
                    if not fallback_text.strip():
                        fallback_text = await asyncio.to_thread(_parse_file_locally, file_ext, file_content)
                    logger.info(
                        "Remote file parsing completed filename=%r duration_seconds=%.1f extracted_chars=%s",
                        filename,
                        time.monotonic() - parse_started_at,
                        len(fallback_text),
                    )
            except httpx.HTTPStatusError as exc:
                logger.warning(
                    "Remote file parser returned HTTP %s after %.1fs: %s; falling back to local parsing",
                    exc.response.status_code,
                    time.monotonic() - parse_started_at,
                    exc.response.text[:1000],
                )
                fallback_text = await asyncio.to_thread(_parse_file_locally, file_ext, file_content)
            except (httpx.HTTPError, ValueError) as exc:
                logger.warning(
                    "Remote file parser failed with %s after %.1fs: %s; falling back to local parsing",
                    type(exc).__name__,
                    time.monotonic() - parse_started_at,
                    exc,
                )
                fallback_text = await asyncio.to_thread(_parse_file_locally, file_ext, file_content)

        attachment_path = os.path.join(ATTACHMENT_DIR, f"{uuid.uuid4()}_{os.path.basename(filename)}")
        with open(attachment_path, "wb") as output:
            output.write(file_content)
        db = SessionLocal()
        try:
            conversation = _ensure_local_conversation(db, conversation_id, identity=identity)
            attachment = AttachmentModel(
                conversation_id=conversation.id,
                file_name=filename,
                mime_type=file.content_type or "application/octet-stream",
                size_bytes=len(file_content),
                storage_path=attachment_path,
                extracted_text=fallback_text,
            )
            db.add(attachment)
            db.commit()
            ThreadRecorder(db, conversation.id).record("attachment_added", {
                "attachment_id": attachment.id,
                "file_name": filename,
                "mime_type": attachment.mime_type,
                "size_bytes": attachment.size_bytes,
            })
            res_json = {"results": {filename: {"md_content": fallback_text}}}
            return JSONResponse(content={
                "success": True,
                "attachment_id": attachment.id,
                "extracted_text": fallback_text,
                "result": res_json,
            })
        finally:
            db.close()

    except HTTPException:
        raise
    except Exception as e:
        logger.exception("File processing failed filename=%r conversation_id=%s", file.filename, conversation_id)
        raise HTTPException(status_code=500, detail="File processing failed") from e

@app.post("/chat/stream")
async def chat_stream(request: ChatRequest, identity: RequestIdentity = Depends(get_request_identity)):
    from agent import run_agent
    identity = normalize_identity(identity)
    registry = _accessible_skill_registry(identity, load_instructions=False)
    invalid_skills = [skill_id for skill_id in request.selected_skill_ids if registry.get_skill(skill_id) is None]
    if invalid_skills:
        raise HTTPException(status_code=422, detail=f"所选技能不存在或无权访问: {', '.join(invalid_skills)}")

    async def event_generator():
        async for chunk in run_agent(
            query=request.query,
            context_text=request.context_text,
            conversation_id=request.conversation_id,
            attachment_name=request.attachment_name,
            attachment_id=request.attachment_id,
            attachment_ids=request.attachment_ids,
            selected_skill_ids=request.selected_skill_ids,
            identity=identity,
        ):
            yield chunk

    return StreamingResponse(
        event_generator(),
        media_type="text/event-stream",
        headers={
            "Cache-Control": "no-cache, no-transform",
            "Connection": "keep-alive",
            "X-Accel-Buffering": "no",
        },
    )


@app.post("/api/v1/runs/{run_id}/retry")
async def retry_run(run_id: str, identity: RequestIdentity = Depends(get_request_identity)):
    from agent import run_agent

    identity = normalize_identity(identity)
    db = SessionLocal()
    try:
        source_run = owned_run(db, run_id, identity)
        if source_run is None:
            raise HTTPException(status_code=404, detail="Run not found")
        if source_run.status in {"pending", "running", "awaiting_approval"}:
            raise HTTPException(status_code=409, detail="Run is still active")

        user_message = (
            db.query(MessageModel)
            .filter(MessageModel.run_id == run_id, MessageModel.role == "user")
            .order_by(MessageModel.created_at, MessageModel.id)
            .first()
        )
        run_started = (
            db.query(RunEventModel)
            .filter(RunEventModel.run_id == run_id, RunEventModel.event_type == "run_started")
            .order_by(RunEventModel.sequence)
            .first()
        )
        user_event = (
            db.query(ThreadEventModel)
            .filter(ThreadEventModel.run_id == run_id, ThreadEventModel.event_type == "user_message")
            .order_by(ThreadEventModel.sequence)
            .first()
        )
        run_payload = run_started.payload or {} if run_started else {}
        user_payload = user_event.payload or {} if user_event else {}
        query = (user_message.content if user_message else None) or str(run_payload.get("query") or "").strip()
        if not query:
            raise HTTPException(status_code=409, detail="Original user request is unavailable")
        attachment_ids = list(dict.fromkeys(user_payload.get("attachment_ids") or []))
        if not attachment_ids:
            attachment_ids = [
                attachment.id
                for attachment in db.query(AttachmentModel).filter(AttachmentModel.run_id == run_id).all()
            ]
        attachment_name = (
            user_message.attachment_name if user_message else None
        ) or user_payload.get("attachment_name")
        conversation_id = source_run.conversation_id
        recorded_requested_skills = user_payload.get("requested_skill_ids")
        if isinstance(recorded_requested_skills, list):
            selected_skill_ids = list(dict.fromkeys(recorded_requested_skills))[:1]
        else:
            # Older runs stored user-selected and automatically recommended Skills
            # together. Preserve the first preference without violating single-select.
            selected_skill_ids = list(dict.fromkeys(source_run.selected_skills or []))[:1]
    finally:
        db.close()

    async def event_generator():
        async for chunk in run_agent(
            query=query,
            conversation_id=conversation_id,
            attachment_name=attachment_name,
            attachment_ids=attachment_ids,
            selected_skill_ids=selected_skill_ids,
            identity=identity,
            persist_user_message=False,
            retry_of_run_id=run_id,
        ):
            yield chunk

    logger.info("Retrying run source_run_id=%s conversation_id=%s", run_id, conversation_id)
    return StreamingResponse(
        event_generator(),
        media_type="text/event-stream",
        headers={
            "Cache-Control": "no-cache, no-transform",
            "Connection": "keep-alive",
            "X-Accel-Buffering": "no",
        },
    )


@app.get("/api/v1/conversations")
async def list_conversations(
    include_archived: bool = False,
    identity: RequestIdentity = Depends(get_request_identity),
):
    db = SessionLocal()
    try:
        identity = normalize_identity(identity)
        user = db.query(UserModel).filter(
            UserModel.tenant_id == identity.tenant_id,
            UserModel.external_id == identity.user_id,
        ).first()
        if user is None:
            return []
        last_content = (
            select(MessageModel.content)
            .where(MessageModel.conversation_id == ConversationModel.id)
            .order_by(MessageModel.created_at.desc(), MessageModel.role.desc(), MessageModel.id.desc())
            .limit(1)
            .correlate(ConversationModel)
            .scalar_subquery()
        )
        last_role = (
            select(MessageModel.role)
            .where(MessageModel.conversation_id == ConversationModel.id)
            .order_by(MessageModel.created_at.desc(), MessageModel.role.desc(), MessageModel.id.desc())
            .limit(1)
            .correlate(ConversationModel)
            .scalar_subquery()
        )
        conversations = (
            db.query(
                ConversationModel,
                ConversationStateModel.is_archived,
                ConversationStateModel.is_pinned,
                last_content.label("last_message_content"),
                last_role.label("last_message_role"),
            )
            .outerjoin(ConversationStateModel, ConversationStateModel.conversation_id == ConversationModel.id)
            .filter(ConversationModel.user_id == user.id)
            .order_by(
                func.coalesce(ConversationStateModel.is_pinned, False).desc(),
                ConversationModel.updated_at.desc(),
            )
            .all()
        )
        result = []
        for conversation, is_archived, is_pinned, last_message_content, last_message_role in conversations:
            is_archived = bool(is_archived)
            if is_archived and not include_archived:
                continue
            last_message_content = last_message_content or ""
            if last_message_role == "assistant":
                _process_preamble, last_message_content = split_leading_process_preamble(last_message_content)
            result.append({
                "id": conversation.id,
                "title": conversation.title,
                "created_at": conversation.created_at,
                "updated_at": conversation.updated_at,
                "last_message": last_message_content[:80],
                "is_archived": is_archived,
                "is_pinned": bool(is_pinned),
            })
        return result
    finally:
        db.close()


@app.get("/api/v1/conversations/{conversation_id}")
async def get_conversation(
    conversation_id: str,
    identity: RequestIdentity = Depends(get_request_identity),
):
    db = SessionLocal()
    try:
        conversation = owned_conversation(db, conversation_id, normalize_identity(identity))
        if conversation is None:
            raise HTTPException(status_code=404, detail="Conversation not found")
        messages = (
            db.query(MessageModel)
            .filter(MessageModel.conversation_id == conversation_id)
            .order_by(MessageModel.created_at, MessageModel.role.desc(), MessageModel.id)
            .all()
        )
        run_ids = {message.run_id for message in messages if message.run_id}
        runs_by_id = {
            run.id: run
            for run in db.query(RunModel).filter(RunModel.id.in_(run_ids)).all()
        } if run_ids else {}
        artifacts_by_run = {run_id: [] for run_id in run_ids}
        if run_ids:
            for artifact in (
                db.query(ArtifactModel)
                .filter(ArtifactModel.run_id.in_(run_ids))
                .order_by(ArtifactModel.created_at)
                .all()
            ):
                artifacts_by_run[artifact.run_id].append(artifact)
        events_by_run = {run_id: [] for run_id in run_ids}
        attachment_names_by_run = {run_id: [] for run_id in run_ids}
        if run_ids:
            for attachment in db.query(AttachmentModel).filter(AttachmentModel.run_id.in_(run_ids)).order_by(AttachmentModel.created_at).all():
                attachment_names_by_run[attachment.run_id].append(attachment.file_name)
            for run_event in (
                db.query(RunEventModel)
                .filter(RunEventModel.run_id.in_(run_ids))
                .order_by(RunEventModel.run_id, RunEventModel.sequence)
                .all()
            ):
                events_by_run[run_event.run_id].append(run_event)
        superseded_run_ids = {
            str(run_event.payload.get("retry_of_run_id"))
            for run_events in events_by_run.values()
            for run_event in run_events
            if run_event.event_type == "run_started"
            and isinstance(run_event.payload, dict)
            and run_event.payload.get("retry_of_run_id")
        }
        serialized_messages = []
        for message in messages:
            if message.role == "assistant" and message.run_id in superseded_run_ids:
                continue
            artifacts = []
            message_run = runs_by_id.get(message.run_id) if message.run_id else None
            process_completed_at = None
            if message_run:
                process_completed_at = message_run.completed_at
                if process_completed_at is None and message_run.status in {"completed", "failed", "cancelled"}:
                    process_completed_at = message.created_at
            if message.role == "assistant" and message.run_id:
                stored_artifacts = artifacts_by_run.get(message.run_id, [])
                for artifact in stored_artifacts:
                    preview_kind, preview = read_artifact_preview(artifact)
                    artifacts.append({
                        "artifact_id": artifact.id,
                        "artifact_type": artifact.artifact_type,
                        "title": artifact.title,
                        "preview_kind": preview_kind,
                        preview_kind: preview,
                        "download_url": create_download_url(artifact.id),
                        "preview_url": create_preview_url(artifact.id),
                        "missing_fields": artifact.missing_fields,
                    })
            content = message.content
            if message.role == "assistant" and contains_artifact_source(content):
                content = sanitize_artifact_final_text(content) if artifacts else "该轮任务未生成有效产物，请重新执行。"
            if message.role == "assistant":
                _process_preamble, content = split_leading_process_preamble(content)
            serialized_messages.append({
                "id": message.id,
                "role": message.role,
                "content": content,
                "attachment_name": message.attachment_name,
                "attachment_names": attachment_names_by_run.get(message.run_id, []) if message.run_id else [],
                "created_at": message.created_at,
                "artifacts": artifacts,
                "groups": _groups_from_run_events(message.run_id, events_by_run.get(message.run_id, [])) if message.role == "assistant" and message.run_id else [],
                "process": ({
                    "run_id": message.run_id,
                    "status": message_run.status,
                    "started_at": message_run.created_at.isoformat(),
                    "completed_at": process_completed_at.isoformat() if process_completed_at else None,
                } if message.role == "assistant" and message_run else None),
            })
        pending_run = (
            db.query(RunModel)
            .filter(
                RunModel.conversation_id == conversation_id,
                RunModel.status == "awaiting_approval",
            )
            .order_by(RunModel.created_at.desc())
            .first()
        )
        pending_approval = None
        if pending_run and pending_run.pending_approval:
            pending_approval = {"run_id": pending_run.id, **pending_run.pending_approval}
        return {
            "id": conversation.id,
            "title": conversation.title,
            "created_at": conversation.created_at,
            "updated_at": conversation.updated_at,
            "is_archived": get_or_create_state(db, conversation_id).is_archived,
            "is_pinned": get_or_create_state(db, conversation_id).is_pinned,
            "pending_approval": pending_approval,
            "messages": serialized_messages,
        }
    finally:
        db.close()


@app.get("/api/v1/conversations/{conversation_id}/events")
async def get_conversation_events(
    conversation_id: str,
    after: int = 0,
    identity: RequestIdentity = Depends(get_request_identity),
):
    db = SessionLocal()
    try:
        if owned_conversation(db, conversation_id, normalize_identity(identity)) is None:
            raise HTTPException(status_code=404, detail="Conversation not found")
        events = ThreadRecorder(db, conversation_id).events_after(max(0, after))
        return {
            "conversation_id": conversation_id,
            "events": [{
                "sequence": event.sequence,
                "type": event.event_type,
                "run_id": event.run_id,
                "payload": event.payload,
                "created_at": event.created_at,
            } for event in events],
        }
    finally:
        db.close()


@app.post("/api/v1/conversations/{conversation_id}/state")
async def set_conversation_state(
    conversation_id: str,
    state: ConversationStateRequest,
    identity: RequestIdentity = Depends(get_request_identity),
):
    db = SessionLocal()
    try:
        conversation = owned_conversation(db, conversation_id, normalize_identity(identity))
        if conversation is None:
            raise HTTPException(status_code=404, detail="Conversation not found")
        result = update_state(
            db,
            conversation_id,
            is_archived=state.is_archived,
            is_pinned=state.is_pinned,
        )
        ThreadRecorder(db, conversation_id).record("thread_state_changed", {
            "is_archived": result.is_archived,
            "is_pinned": result.is_pinned,
        })
        return {"id": conversation_id, "is_archived": result.is_archived, "is_pinned": result.is_pinned}
    finally:
        db.close()


@app.patch("/api/v1/conversations/{conversation_id}")
async def rename_conversation(
    conversation_id: str,
    request: ConversationUpdateRequest,
    identity: RequestIdentity = Depends(get_request_identity),
):
    db = SessionLocal()
    try:
        conversation = owned_conversation(db, conversation_id, normalize_identity(identity))
        if conversation is None:
            raise HTTPException(status_code=404, detail="Conversation not found")
        conversation.title = request.title.strip()
        conversation.updated_at = utc_now()
        ThreadRecorder(db, conversation_id).record("thread_renamed", {"title": conversation.title})
        db.commit()
        return {"id": conversation_id, "title": conversation.title}
    finally:
        db.close()


@app.delete("/api/v1/conversations/{conversation_id}")
async def delete_conversation(
    conversation_id: str,
    identity: RequestIdentity = Depends(get_request_identity),
):
    db = SessionLocal()
    try:
        conversation = owned_conversation(db, conversation_id, normalize_identity(identity))
        if conversation is None:
            raise HTTPException(status_code=404, detail="Conversation not found")
        storage_paths = [attachment.storage_path for attachment in conversation.attachments]
        for run in conversation.runs:
            storage_paths.extend(artifact.storage_path for artifact in run.artifacts)
        db.delete(conversation)
        db.commit()
        for storage_path in storage_paths:
            if not storage_path:
                continue
            try:
                if os.path.isdir(storage_path):
                    shutil.rmtree(storage_path)
                elif os.path.exists(storage_path):
                    os.remove(storage_path)
                preview_path = f"{storage_path}.preview.md"
                if os.path.exists(preview_path):
                    os.remove(preview_path)
            except OSError:
                logger.warning("Failed to remove conversation file: %s", storage_path, exc_info=True)
        return {"id": conversation_id, "deleted": True}
    finally:
        db.close()


@app.get("/api/v1/attachments/{attachment_id}/download")
async def download_attachment(
    attachment_id: str,
    identity: RequestIdentity = Depends(get_request_identity),
):
    db = SessionLocal()
    try:
        attachment = owned_attachment(db, attachment_id, normalize_identity(identity))
        if attachment is None or not os.path.exists(attachment.storage_path):
            raise HTTPException(status_code=404, detail="Attachment not found")
        return FileResponse(
            path=attachment.storage_path,
            filename=attachment.file_name,
            media_type=attachment.mime_type,
        )
    finally:
        db.close()

@app.get("/api/v1/artifacts/{artifact_id}/download")
async def download_artifact(
    artifact_id: str,
    token: str,
    expires: int,
    identity: RequestIdentity = Depends(get_request_identity),
):
    if not verify_download_token(artifact_id, token, expires):
        raise HTTPException(status_code=403, detail="Invalid download token")
    db = SessionLocal()
    try:
        art = owned_artifact(db, artifact_id, normalize_identity(identity))
        if not art or not os.path.exists(art.storage_path):
            raise HTTPException(status_code=404, detail="Artifact file not found")
        
        filename = _artifact_download_filename(art.title, art.storage_path)
        return FileResponse(path=art.storage_path, filename=filename, media_type=art.mime_type)
    finally:
        db.close()


@app.get("/api/v1/artifacts/{artifact_id}/preview")
async def preview_artifact(
    artifact_id: str,
    token: str,
    expires: int,
    identity: RequestIdentity = Depends(get_request_identity),
):
    if not verify_download_token(artifact_id, token, expires):
        raise HTTPException(status_code=403, detail="Invalid preview token")
    db = SessionLocal()
    try:
        art = owned_artifact(db, artifact_id, normalize_identity(identity))
        if not art or not os.path.exists(art.storage_path):
            raise HTTPException(status_code=404, detail="Artifact file not found")
        preview_path = art.storage_path
        preview_mime_type = art.mime_type
        if art.mime_type == PPTX_MIME:
            try:
                preview_path = await asyncio.to_thread(ensure_pptx_pdf_preview, art.storage_path)
            except (OSError, RuntimeError, subprocess.SubprocessError, ValueError) as exc:
                logger.exception("Failed to render PPTX preview artifact=%s", artifact_id)
                raise HTTPException(status_code=503, detail=f"PPTX preview conversion failed: {exc}") from exc
            if preview_path is None:
                raise HTTPException(status_code=503, detail="PPTX preview converter is not configured")
            preview_mime_type = "application/pdf"
        elif art.mime_type not in {"application/pdf"} and not art.mime_type.startswith("image/"):
            raise HTTPException(status_code=415, detail="This artifact type has no native browser preview")
        return FileResponse(
            path=preview_path,
            media_type=preview_mime_type,
            headers={"Content-Disposition": "inline"},
        )
    finally:
        db.close()


@app.get("/api/v1/runs/{run_id}")
async def get_run(run_id: str, identity: RequestIdentity = Depends(get_request_identity)):
    db = SessionLocal()
    try:
        run = owned_run(db, run_id, normalize_identity(identity))
        if run is None:
            raise HTTPException(status_code=404, detail="Run not found")
        events = (
            db.query(RunEventModel)
            .filter(RunEventModel.run_id == run_id)
            .order_by(RunEventModel.sequence)
            .all()
        )
        return {
            "id": run.id,
            "conversation_id": run.conversation_id,
            "intent_type": run.intent_type,
            "engine": run.engine,
            "router_confidence": run.router_confidence,
            "router_reasons": run.router_reasons or [],
            "status": run.status,
            "pending_approval": run.pending_approval,
            "error_message": run.error_message,
            "created_at": run.created_at,
            "completed_at": run.completed_at,
            "events": [
                {
                    "sequence": event.sequence,
                    "type": event.event_type,
                    "payload": event.payload,
                    "created_at": event.created_at,
                }
                for event in events
            ],
            "plan_tasks": [{
                "id": task.id,
                "position": task.position,
                "title": task.title,
                "status": task.status,
                "depends_on": task.depends_on or [],
                "attempt": task.attempt,
                "error_message": task.error_message,
            } for task in db.query(PlanTaskModel).filter(PlanTaskModel.run_id == run_id).order_by(PlanTaskModel.position).all()],
        }
    finally:
        db.close()


@app.post("/api/v1/runs/{run_id}/cancel")
async def cancel_run(run_id: str, identity: RequestIdentity = Depends(get_request_identity)):
    db = SessionLocal()
    try:
        run = owned_run(db, run_id, normalize_identity(identity))
        if run is None:
            raise HTTPException(status_code=404, detail="Run not found")
        if run.status not in {"pending", "running", "awaiting_approval"}:
            raise HTTPException(status_code=409, detail=f"Run is already {run.status}")
        run.status = "cancelled"
        run.completed_at = utc_now()
        db.commit()
        return {"id": run.id, "status": run.status}
    finally:
        db.close()


@app.post("/api/v1/runs/{run_id}/approval")
async def decide_run_approval(
    run_id: str,
    request: ApprovalRequest,
    identity: RequestIdentity = Depends(get_request_identity),
):
    db = SessionLocal()
    try:
        identity = normalize_identity(identity)
        run = owned_run(db, run_id, identity)
        if run is None:
            raise HTTPException(status_code=404, detail="Run not found")
        if run.status != "awaiting_approval" or not run.pending_approval:
            raise HTTPException(status_code=409, detail="Run is not awaiting approval")
    finally:
        db.close()
    from agent import resume_agent
    response_text = (request.response or "").strip()
    if request.decision is None and not response_text:
        raise HTTPException(status_code=422, detail="请通过对话回复是否同意该方案")
    decision = request.decision or _approval_decision_from_response(response_text)
    return StreamingResponse(
        resume_agent(run_id, decision, identity, user_response=response_text),
        media_type="text/event-stream",
        headers={"Cache-Control": "no-cache, no-transform", "X-Accel-Buffering": "no"},
    )

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(
        app,
        host=os.getenv("APP_HOST", "127.0.0.1"),
        port=int(os.getenv("APP_PORT", "14499")),
    )
