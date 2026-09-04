import asyncio
import io
import json
import sys
import tempfile
import time
import unittest
import uuid
import zipfile
from pathlib import Path
from unittest.mock import AsyncMock, patch

from sqlalchemy import create_engine, event
from sqlalchemy.orm import sessionmaker
from docx.oxml.ns import qn

from starlette.datastructures import UploadFile


BACKEND_DIR = Path(__file__).resolve().parents[1] / "backend"
sys.path.insert(0, str(BACKEND_DIR))

import agent
import main
from rag.chunker import DocumentChunker
from security import create_download_token, verify_download_token
from capabilities.skill_registry import SkillRegistry
from database import Base
from models import ArtifactModel, AttachmentModel, ConversationModel, ConversationStateModel, MessageModel, PlanTaskModel, RunEventModel, RunModel, ThreadEventModel, UserModel
from exporters.docx_exporter import export_markdown_to_docx


class StartupRecoveryTests(unittest.TestCase):
    def test_startup_finalizes_only_runs_without_a_live_executor(self):
        engine = create_engine("sqlite:///:memory:")
        Base.metadata.create_all(engine)
        session_factory = sessionmaker(bind=engine)
        db = session_factory()
        user = UserModel(username="startup-user")
        db.add(user)
        db.flush()
        conversation = ConversationModel(user_id=user.id, title="启动恢复")
        db.add(conversation)
        db.flush()
        pending = RunModel(conversation_id=conversation.id, status="pending")
        running = RunModel(conversation_id=conversation.id, status="running")
        approval = RunModel(conversation_id=conversation.id, status="awaiting_approval")
        db.add_all([pending, running, approval])
        db.flush()
        db.add_all([
            PlanTaskModel(run_id=pending.id, position=1, title="等待执行", status="pending"),
            PlanTaskModel(run_id=running.id, position=1, title="执行中", status="running"),
            PlanTaskModel(run_id=approval.id, position=1, title="等待确认", status="pending"),
        ])
        db.commit()
        pending_id, running_id, approval_id = pending.id, running.id, approval.id
        db.close()

        with patch.object(main, "SessionLocal", session_factory):
            count = main._finalize_interrupted_runs()

        db = session_factory()
        self.assertEqual(count, 2)
        self.assertEqual(db.get(RunModel, pending_id).status, "failed")
        self.assertEqual(db.get(RunModel, running_id).status, "failed")
        self.assertEqual(db.get(RunModel, approval_id).status, "awaiting_approval")
        self.assertEqual(db.query(PlanTaskModel).filter_by(run_id=running_id).one().status, "failed")
        self.assertEqual(db.query(PlanTaskModel).filter_by(run_id=approval_id).one().status, "pending")
        db.close()


class DetachedEventStreamTests(unittest.IsolatedAsyncioTestCase):
    async def asyncTearDown(self):
        if main._BACKGROUND_RUN_TASKS:
            await asyncio.gather(*tuple(main._BACKGROUND_RUN_TASKS), return_exceptions=True)

    async def test_stream_sends_heartbeat_while_agent_is_idle(self):
        async def source():
            await asyncio.sleep(0.03)
            yield "data: done\n\n"

        stream = main._detached_event_stream(source(), heartbeat_seconds=0.01)
        self.assertEqual(await anext(stream), ": keep-alive\n\n")
        item = await anext(stream)
        while item.startswith(":"):
            item = await anext(stream)
        self.assertEqual(item, "data: done\n\n")
        await stream.aclose()

    async def test_agent_continues_after_http_consumer_disconnects(self):
        completed = asyncio.Event()

        async def source():
            yield "data: started\n\n"
            await asyncio.sleep(0.02)
            completed.set()
            yield "data: completed\n\n"

        stream = main._detached_event_stream(source(), heartbeat_seconds=1)
        self.assertEqual(await anext(stream), "data: started\n\n")
        await stream.aclose()
        await asyncio.wait_for(completed.wait(), timeout=0.2)

class UploadTests(unittest.IsolatedAsyncioTestCase):
    def test_default_cors_supports_localhost_and_loopback(self):
        self.assertIn("http://localhost:6080", main.ALLOWED_ORIGINS)
        self.assertIn("http://127.0.0.1:6080", main.ALLOWED_ORIGINS)

    def test_local_pdf_parser_extracts_text(self):
        import pymupdf

        document = pymupdf.open()
        page = document.new_page()
        page.insert_text((72, 72), "PDF parser works")
        content = document.tobytes()
        document.close()

        self.assertIn("PDF parser works", main._parse_file_locally("pdf", content))

    async def test_root_redirects_to_frontend(self):
        response = await main.frontend_entry()
        self.assertEqual(response.status_code, 307)
        self.assertEqual(response.headers["location"], main.FRONTEND_URL)

    async def test_unsupported_extension_keeps_415_status(self):
        upload = UploadFile(file=io.BytesIO(b"data"), filename="payload.exe")
        with self.assertRaises(main.HTTPException) as raised:
            await main.upload_file(upload)
        self.assertEqual(raised.exception.status_code, 415)

    async def test_upload_rejects_file_above_configured_limit(self):
        upload = UploadFile(file=io.BytesIO(b"1234"), filename="brief.txt")
        with patch.object(main, "MAX_UPLOAD_BYTES", 3):
            with self.assertRaises(main.HTTPException) as raised:
                await main.upload_file(upload)
        self.assertEqual(raised.exception.status_code, 413)

    async def test_skill_zip_is_preflighted_before_private_install(self):
        content = io.BytesIO()
        with zipfile.ZipFile(content, "w") as archive:
            archive.writestr("user-skill/SKILL.md", "# User skill")
        upload = UploadFile(file=io.BytesIO(content.getvalue()), filename="user-skill.zip")
        with (
            tempfile.TemporaryDirectory() as directory,
            patch.object(main, "get_user_skill_root", return_value=Path(directory)),
            patch.object(main, "_generate_skill_intro", new=AsyncMock(return_value="用户技能简介")),
        ):
            result = await main.upload_skill(
                file=upload,
                identity=agent.LOCAL_IDENTITY,
            )
            self.assertTrue(result["success"])
            self.assertTrue((Path(directory) / "user-skill" / "SKILL.md").is_file())

    async def test_skill_preflight_failure_is_returned_and_not_installed(self):
        content = io.BytesIO()
        with zipfile.ZipFile(content, "w") as archive:
            archive.writestr("broken-skill/readme.md", "missing entrypoint")
        upload = UploadFile(file=io.BytesIO(content.getvalue()), filename="broken-skill.zip")
        with tempfile.TemporaryDirectory() as directory, patch.object(
            main, "get_user_skill_root", return_value=Path(directory)
        ):
            with self.assertRaises(main.HTTPException) as raised:
                await main.upload_skill(
                    file=upload,
                    identity=agent.LOCAL_IDENTITY,
                )
            self.assertEqual(raised.exception.status_code, 422)
            self.assertIn("预检未通过", raised.exception.detail)
            self.assertEqual(list(Path(directory).iterdir()), [])

    async def test_remote_failure_uses_local_text_parser(self):
        class FailingClient:
            async def __aenter__(self):
                return self

            async def __aexit__(self, *_args):
                return None

            async def post(self, *_args, **_kwargs):
                raise main.httpx.HTTPError("offline")

        upload = UploadFile(file=io.BytesIO("有效文本".encode()), filename="brief.txt")
        engine = create_engine("sqlite:///:memory:")
        Base.metadata.create_all(engine)
        session_factory = sessionmaker(bind=engine)
        with (
            patch.object(main, "FILE_PARSE_URL", "https://parser.invalid"),
            patch.object(main, "FILE_PARSE_TOKEN", "token"),
            patch.object(main.httpx, "AsyncClient", return_value=FailingClient()),
            patch.object(main, "SessionLocal", session_factory),
        ):
            response = await main.upload_file(upload)

        payload = json.loads(response.body)
        self.assertEqual(payload["extracted_text"], "有效文本")
        self.assertEqual(payload["result"]["results"]["brief.txt"]["md_content"], "有效文本")

    async def test_remote_parser_uses_configured_pipeline_and_timeout(self):
        captured = {}

        class SuccessfulResponse:
            def raise_for_status(self):
                return None

            def json(self):
                return {"results": {"brief.txt": {"md_content": "远程解析结果"}}}

        class CapturingClient:
            def __init__(self, timeout):
                captured["timeout"] = timeout

            async def __aenter__(self):
                return self

            async def __aexit__(self, *_args):
                return None

            async def post(self, *_args, **kwargs):
                captured["data"] = kwargs["data"]
                return SuccessfulResponse()

        upload = UploadFile(file=io.BytesIO("本地内容".encode()), filename="brief.txt")
        engine = create_engine("sqlite:///:memory:")
        Base.metadata.create_all(engine)
        session_factory = sessionmaker(bind=engine)
        with (
            patch.object(main, "FILE_PARSE_URL", "https://parser.example/file_parse"),
            patch.object(main, "FILE_PARSE_TOKEN", "token"),
            patch.object(main, "FILE_PARSE_BACKEND", "pipeline"),
            patch.object(main, "FILE_PARSE_TIMEOUT_SECONDS", 600.0),
            patch.object(main.httpx, "AsyncClient", CapturingClient),
            patch.object(main, "SessionLocal", session_factory),
        ):
            response = await main.upload_file(upload)

        payload = json.loads(response.body)
        self.assertEqual(payload["extracted_text"], "远程解析结果")
        self.assertEqual(captured["timeout"], 600.0)
        self.assertEqual(captured["data"]["backend"], "pipeline")
        self.assertEqual(captured["data"]["lang_list"], "ch")
        self.assertEqual(captured["data"]["parse_method"], "auto")

    def test_docx_export_preserves_chinese_and_markdown_content(self):
        import tempfile
        from docx import Document

        content = "# 总结\n正文含 **重点内容**。\n- 条目一\n1. 条目二\n\n| 项目 | 状态 |\n| --- | --- |\n| 交付 | 完成 |"
        with tempfile.TemporaryDirectory() as temp_dir:
            output = Path(temp_dir) / "report.docx"
            export_markdown_to_docx(content, str(output), title="测试文档")
            document = Document(output)

        full_text = "\n".join(paragraph.text for paragraph in document.paragraphs)
        self.assertIn("重点内容", full_text)
        self.assertIn("条目一", full_text)
        self.assertEqual(document.styles["Normal"].font.name, "宋体")
        self.assertEqual(document.styles["Normal"]._element.rPr.rFonts.get(qn("w:eastAsia")), "宋体")
        self.assertEqual(document.tables[0].cell(1, 0).text, "交付")

    def test_local_docx_parser_keeps_table_cells(self):
        from docx import Document

        document = Document()
        document.add_paragraph("正文段落")
        table = document.add_table(rows=1, cols=2)
        table.cell(0, 0).text = "项目"
        table.cell(0, 1).text = "状态"
        stream = io.BytesIO()
        document.save(stream)

        extracted = main._parse_file_locally("docx", stream.getvalue())
        self.assertIn("正文段落", extracted)
        self.assertIn("| 项目 | 状态 |", extracted)


class RoutingAndRetrievalTests(unittest.TestCase):
    def test_llm_base_url_collapses_duplicate_path_slashes(self):
        normalized = agent._normalize_base_url("http://gateway.example//v1/")
        self.assertEqual(normalized, "http://gateway.example/v1")

    def test_dialogue_approval_understands_affirmation_and_feedback(self):
        for response in ("好的", "可以，继续执行", "同意上述方案", "没问题，开始吧"):
            self.assertEqual(main._approval_decision_from_response(response), "approve")
        for response in ("先别执行", "可以修改第二点吗？", "不同意", "不行", "把配色调整一下"):
            self.assertEqual(main._approval_decision_from_response(response), "reject")

    def test_general_agent_has_no_fixed_intent_router(self):
        self.assertFalse(hasattr(agent, "IntentRouter"))
        self.assertIn("普通问答", agent._base_system_prompt(""))

    def test_report_uses_report_skill(self):
        self.assertEqual(SkillRegistry().get_skill("report").name, "report")

    def test_registry_recommends_skills_from_general_metadata(self):
        registry = SkillRegistry()
        self.assertEqual(registry.recommend_skills("请生成一份旅游推介PPT")[0].name, "ppt")
        self.assertEqual(registry.recommend_skills("编制交通项目投标文件")[0].name, "bidding")
        self.assertEqual(registry.recommend_skills("生成一份行业调研报告")[0].name, "report")
        self.assertEqual(registry.recommend_skills("撰写项目实施方案")[0].name, "document")
        self.assertEqual(registry.recommend_skills("你好，今天天气怎么样"), [])

    def test_explicit_skill_stays_excluded_from_additional_matches(self):
        registry = SkillRegistry()
        recommended = registry.recommend_skills("把调研报告做成PPT", exclude={"report"})
        self.assertNotIn("report", [skill.name for skill in recommended])
        self.assertIn("ppt", [skill.name for skill in recommended])

    def test_chinese_query_retrieves_relevant_chunk(self):
        chunks = [
            {"content": "开篇无关内容", "token_count": 10},
            {"content": "项目质量与安全措施详述", "token_count": 10},
        ]
        result = DocumentChunker.retrieve_relevant_chunks(
            chunks, "请分析质量与安全措施", top_k=1
        )
        self.assertEqual(result, ["项目质量与安全措施详述"])

    def test_download_tokens_are_artifact_specific(self):
        expires = int(time.time()) + 60
        token = create_download_token("artifact-a", expires)
        self.assertTrue(verify_download_token("artifact-a", token, expires))
        self.assertFalse(verify_download_token("artifact-b", token, expires))
        self.assertFalse(verify_download_token("artifact-a", token, 0))

    def test_artifact_download_filename_uses_model_title_and_storage_extension(self):
        self.assertEqual(
            main._artifact_download_filename("季度汇报", "storage/run-1.pptx"),
            "季度汇报.pptx",
        )
        self.assertEqual(
            main._artifact_download_filename("季度/汇报.pptx", "storage/run-1.pptx"),
            "季度_汇报.pptx",
        )

    def test_binary_sandbox_artifact_is_persisted_without_text_conversion(self):
        payload = b"PK\x03\x04pptx-data"

        class FakeDb:
            def add(self, _artifact):
                return None

            def commit(self):
                return None

            def refresh(self, artifact):
                artifact.id = "artifact-test"

        with tempfile.TemporaryDirectory() as temp_dir, patch.object(agent, "STORAGE_DIR", temp_dir):
            artifact, _result = agent._persist_artifact(FakeDb(), "run-test", {
                "artifact_type": "ppt",
                "title": "沙箱演示文稿",
                "extension": "pptx",
                "mime_type": "application/vnd.openxmlformats-officedocument.presentationml.presentation",
                "content_bytes": payload,
                "preview_kind": "none",
            }, 0)

            self.assertEqual(Path(artifact.storage_path).read_bytes(), payload)

    def test_html_sandbox_artifact_is_available_to_streaming_preview(self):
        payload = "<!doctype html><html><body><h1>实时预览</h1></body></html>".encode("utf-8")

        class FakeDb:
            def add(self, _artifact):
                return None

            def commit(self):
                return None

            def refresh(self, artifact):
                artifact.id = "artifact-html"

        with tempfile.TemporaryDirectory() as temp_dir, patch.object(agent, "STORAGE_DIR", temp_dir):
            _artifact, result = agent._persist_artifact(FakeDb(), "run-html", {
                "artifact_type": "html",
                "title": "交互页面",
                "extension": "html",
                "mime_type": "text/html",
                "content_bytes": payload,
                "preview_kind": "none",
            }, 0)

        self.assertEqual(result["preview_kind"], "html")
        self.assertEqual(result["html"], payload.decode("utf-8"))


class RunControlTests(unittest.IsolatedAsyncioTestCase):
    async def test_pending_approval_can_only_be_claimed_once(self):
        engine = create_engine("sqlite:///:memory:")
        Base.metadata.create_all(engine)
        session_factory = sessionmaker(bind=engine)
        setup = session_factory()
        user = UserModel(username="approval-user")
        setup.add(user)
        setup.flush()
        conversation = ConversationModel(id=str(uuid.uuid4()), user_id=user.id)
        setup.add(conversation)
        run = RunModel(
            conversation_id=conversation.id,
            status="awaiting_approval",
            pending_approval={"name": "request_user_confirmation"},
        )
        setup.add(run)
        setup.commit()
        run_id = run.id
        setup.close()

        first_db = session_factory()
        second_db = session_factory()
        first_run = first_db.get(RunModel, run_id)
        second_run = second_db.get(RunModel, run_id)
        approval = agent._claim_pending_approval(first_db, first_run)
        self.assertEqual(approval["name"], "request_user_confirmation")
        with self.assertRaisesRegex(RuntimeError, "already decided"):
            agent._claim_pending_approval(second_db, second_run)
        first_db.close()
        second_db.close()

    async def test_running_run_can_be_cancelled_and_queried(self):
        engine = create_engine("sqlite:///:memory:")
        Base.metadata.create_all(engine)
        session_factory = sessionmaker(bind=engine)
        conversation_id = str(uuid.uuid4())

        db = session_factory()
        user = UserModel(username="run-control-user")
        db.add(user)
        db.flush()
        db.add(ConversationModel(id=conversation_id, user_id=user.id))
        run = RunModel(conversation_id=conversation_id, status="running")
        db.add(run)
        db.commit()
        run_id = run.id
        db.close()

        with patch.object(main, "SessionLocal", session_factory):
            cancelled = await main.cancel_run(run_id)
            fetched = await main.get_run(run_id)

        self.assertEqual(cancelled["status"], "cancelled")
        self.assertEqual(fetched["status"], "cancelled")
        self.assertEqual(fetched["events"], [])

    async def test_terminal_run_can_be_retried_without_copying_user_message_or_attachment(self):
        engine = create_engine("sqlite:///:memory:")
        Base.metadata.create_all(engine)
        session_factory = sessionmaker(bind=engine)
        conversation_id = str(uuid.uuid4())
        captured = {}

        db = session_factory()
        user = UserModel(username="local-user")
        db.add(user)
        db.flush()
        db.add(ConversationModel(id=conversation_id, user_id=user.id))
        run = RunModel(
            conversation_id=conversation_id,
            status="failed",
            selected_skills=["ppt", "report"],
        )
        db.add(run)
        db.flush()
        attachment = AttachmentModel(
            conversation_id=conversation_id,
            run_id=run.id,
            file_name="brief.txt",
            mime_type="text/plain",
            size_bytes=5,
            storage_path="unused.txt",
            extracted_text="brief",
        )
        db.add(attachment)
        db.flush()
        db.add_all([
            MessageModel(conversation_id=conversation_id, run_id=run.id, role="user", content="重新分析附件", attachment_name="brief.txt"),
            RunEventModel(run_id=run.id, sequence=1, event_type="run_started", payload={"query": "重新分析附件"}),
            ThreadEventModel(
                conversation_id=conversation_id,
                run_id=run.id,
                sequence=1,
                event_type="user_message",
                payload={"attachment_ids": [attachment.id], "attachment_name": "brief.txt"},
            ),
        ])
        db.commit()
        run_id = run.id
        attachment_id = attachment.id
        db.close()

        async def fake_run_agent(**kwargs):
            captured.update(kwargs)
            yield 'data: {"type":"run_started","run_id":"retry-run"}\n\n'

        with (
            patch.object(main, "SessionLocal", session_factory),
            patch.object(agent, "run_agent", fake_run_agent),
        ):
            response = await main.retry_run(run_id)
            chunks = [chunk async for chunk in response.body_iterator]

        self.assertEqual(response.status_code, 200)
        self.assertTrue(chunks)
        self.assertEqual(captured["query"], "重新分析附件")
        self.assertEqual(captured["conversation_id"], conversation_id)
        self.assertEqual(captured["attachment_ids"], [attachment_id])
        self.assertEqual(captured["selected_skill_ids"], ["ppt"])
        self.assertFalse(captured["persist_user_message"])
        self.assertEqual(captured["retry_of_run_id"], run_id)

    async def test_active_run_cannot_be_retried(self):
        engine = create_engine("sqlite:///:memory:")
        Base.metadata.create_all(engine)
        session_factory = sessionmaker(bind=engine)
        conversation_id = str(uuid.uuid4())
        db = session_factory()
        user = UserModel(username="local-user")
        db.add(user)
        db.flush()
        db.add(ConversationModel(id=conversation_id, user_id=user.id))
        run = RunModel(conversation_id=conversation_id, status="running")
        db.add(run)
        db.commit()
        run_id = run.id
        db.close()

        with patch.object(main, "SessionLocal", session_factory):
            with self.assertRaises(main.HTTPException) as raised:
                await main.retry_run(run_id)

        self.assertEqual(raised.exception.status_code, 409)


class ConversationHistoryTests(unittest.IsolatedAsyncioTestCase):
    async def test_conversation_list_is_batched_and_does_not_create_missing_state_rows(self):
        engine = create_engine("sqlite:///:memory:")
        Base.metadata.create_all(engine)
        session_factory = sessionmaker(bind=engine)
        db = session_factory()
        user = UserModel(username="local-user")
        db.add(user)
        db.flush()
        first = ConversationModel(id=str(uuid.uuid4()), user_id=user.id, title="普通任务")
        pinned = ConversationModel(id=str(uuid.uuid4()), user_id=user.id, title="固定任务")
        db.add_all([first, pinned])
        db.flush()
        db.add_all([
            ConversationStateModel(conversation_id=pinned.id, is_pinned=True),
            MessageModel(conversation_id=first.id, role="user", content="第一条消息"),
            MessageModel(conversation_id=pinned.id, role="assistant", content="固定任务结果"),
        ])
        db.commit()
        db.close()

        select_count = 0
        def count_selects(_connection, _cursor, statement, _parameters, _context, _executemany):
            nonlocal select_count
            if statement.lstrip().upper().startswith("SELECT"):
                select_count += 1

        event.listen(engine, "before_cursor_execute", count_selects)
        try:
            with patch.object(main, "SessionLocal", session_factory):
                summaries = await main.list_conversations()
        finally:
            event.remove(engine, "before_cursor_execute", count_selects)

        self.assertEqual([item["title"] for item in summaries], ["固定任务", "普通任务"])
        self.assertEqual(summaries[0]["last_message"], "固定任务结果")
        self.assertFalse(summaries[1]["is_pinned"])
        self.assertLessEqual(select_count, 2)
        db = session_factory()
        self.assertEqual(db.query(ConversationStateModel).count(), 1)
        db.close()

    async def test_conversation_api_hides_corrupted_artifact_source_message(self):
        engine = create_engine("sqlite:///:memory:")
        Base.metadata.create_all(engine)
        session_factory = sessionmaker(bind=engine)
        conversation_id = str(uuid.uuid4())
        db = session_factory()
        user = UserModel(username="local-user")
        db.add(user)
        db.flush()
        db.add(ConversationModel(id=conversation_id, user_id=user.id, title="异常输出"))
        run = RunModel(conversation_id=conversation_id, status="completed")
        db.add(run)
        db.flush()
        db.add(MessageModel(
            conversation_id=conversation_id,
            run_id=run.id,
            role="assistant",
            content="PPT已修改。\n\n[历史产物: 建设方案]\n<!DOCTYPE html><style>body{color:red}</style>",
        ))
        db.commit()
        db.close()

        with patch.object(main, "SessionLocal", session_factory):
            detail = await main.get_conversation(conversation_id)

        self.assertEqual(detail["messages"][0]["content"], "该轮任务未生成有效产物，请重新执行。")

    async def test_conversation_api_replaces_superseded_retry_answer(self):
        engine = create_engine("sqlite:///:memory:")
        Base.metadata.create_all(engine)
        session_factory = sessionmaker(bind=engine)
        conversation_id = str(uuid.uuid4())
        db = session_factory()
        user = UserModel(username="local-user")
        db.add(user)
        db.flush()
        db.add(ConversationModel(id=conversation_id, user_id=user.id, title="覆盖重试"))
        original_run = RunModel(conversation_id=conversation_id, status="completed")
        retry_run = RunModel(conversation_id=conversation_id, status="completed")
        db.add_all([original_run, retry_run])
        db.flush()
        db.add_all([
            MessageModel(conversation_id=conversation_id, run_id=original_run.id, role="user", content="分析问题"),
            MessageModel(conversation_id=conversation_id, run_id=original_run.id, role="assistant", content="旧回答"),
            MessageModel(conversation_id=conversation_id, run_id=retry_run.id, role="assistant", content="新回答"),
            RunEventModel(
                run_id=retry_run.id,
                sequence=1,
                event_type="run_started",
                payload={"retry_of_run_id": original_run.id},
            ),
        ])
        db.commit()
        db.close()

        with patch.object(main, "SessionLocal", session_factory):
            detail = await main.get_conversation(conversation_id)

        self.assertEqual([message["content"] for message in detail["messages"]], ["分析问题", "新回答"])

    async def test_conversation_api_restores_messages_and_artifact_preview(self):
        import tempfile

        engine = create_engine("sqlite:///:memory:")
        Base.metadata.create_all(engine)
        session_factory = sessionmaker(bind=engine)
        conversation_id = str(uuid.uuid4())

        with tempfile.TemporaryDirectory() as temp_dir:
            preview_file = Path(temp_dir) / "deck.html"
            preview_file.write_text("<section class='slide'>历史预览内容足够完整</section>", encoding="utf-8")
            db = session_factory()
            user = UserModel(username="local-user")
            db.add(user)
            db.flush()
            conversation = ConversationModel(id=conversation_id, user_id=user.id, title="历史对话")
            db.add(conversation)
            db.flush()
            run = RunModel(conversation_id=conversation_id, status="completed")
            db.add(run)
            db.flush()
            db.add_all([
                MessageModel(conversation_id=conversation_id, run_id=run.id, role="user", content="生成演示文稿"),
                MessageModel(conversation_id=conversation_id, run_id=run.id, role="assistant", content="已经生成。"),
                RunEventModel(run_id=run.id, sequence=1, event_type="model_started", payload={"turn": 1}),
                RunEventModel(run_id=run.id, sequence=2, event_type="model_delta", payload={"turn": 1, "text": "先读取文档技能说明。"}),
                RunEventModel(run_id=run.id, sequence=3, event_type="model_started", payload={"turn": 2}),
                RunEventModel(run_id=run.id, sequence=4, event_type="model_delta", payload={"turn": 2, "text": "正在整理演示文稿结构。"}),
                RunEventModel(run_id=run.id, sequence=5, event_type="tool_started", payload={"name": "create_document", "arguments": {}}),
                RunEventModel(run_id=run.id, sequence=6, event_type="tool_completed", payload={"name": "create_document"}),
                RunEventModel(run_id=run.id, sequence=7, event_type="model_started", payload={"turn": 3}),
                RunEventModel(run_id=run.id, sequence=8, event_type="model_delta", payload={"turn": 3, "text": "已经生成。"}),
                ArtifactModel(
                    run_id=run.id,
                    title="历史演示文稿",
                    artifact_type="ppt",
                    mime_type="text/html",
                    storage_path=str(preview_file),
                ),
            ])
            db.commit()
            db.close()

            with patch.object(main, "SessionLocal", session_factory):
                summaries = await main.list_conversations()
                detail = await main.get_conversation(conversation_id)

        self.assertEqual(summaries[0]["title"], "历史对话")
        self.assertEqual([item["role"] for item in detail["messages"]], ["user", "assistant"])
        artifact = detail["messages"][1]["artifacts"][0]
        self.assertEqual(artifact["title"], "历史演示文稿")
        self.assertIn("历史预览内容", artifact["html"])
        self.assertIn("/download?", artifact["download_url"])
        self.assertIn("/preview?", artifact["preview_url"])
        groups = detail["messages"][1]["groups"]
        self.assertEqual([group["title"] for group in groups], ["第 1 轮思考", "第 2 轮思考", "正在生成任务产物"])
        self.assertEqual(groups[0]["thoughts"], ["先读取文档技能说明。"])
        self.assertEqual(groups[1]["thoughts"], ["正在整理演示文稿结构。"])
        self.assertEqual(groups[2]["actions"][-1], {"text": "正在执行 create_document", "status": "done"})
        self.assertEqual(detail["messages"][1]["process"]["status"], "completed")

    async def test_conversation_api_moves_leading_process_preamble_to_thoughts(self):
        engine = create_engine("sqlite:///:memory:")
        Base.metadata.create_all(engine)
        session_factory = sessionmaker(bind=engine)
        conversation_id = str(uuid.uuid4())
        db = session_factory()
        user = UserModel(username="local-user")
        db.add(user)
        db.flush()
        db.add(ConversationModel(id=conversation_id, user_id=user.id, title="代码分析"))
        run = RunModel(conversation_id=conversation_id, status="completed")
        db.add(run)
        db.flush()
        raw_content = "我来分析这段代码。\n\n让我先梳理代码逻辑。\n\n## 原因分析\n\n端口未正确释放。"
        db.add_all([
            MessageModel(conversation_id=conversation_id, run_id=run.id, role="assistant", content=raw_content),
            RunEventModel(run_id=run.id, sequence=1, event_type="model_started", payload={"turn": 1}),
            RunEventModel(run_id=run.id, sequence=2, event_type="model_delta", payload={"text": raw_content}),
        ])
        db.commit()
        db.close()

        with patch.object(main, "SessionLocal", session_factory):
            detail = await main.get_conversation(conversation_id)

        message = detail["messages"][0]
        self.assertEqual(message["content"], "## 原因分析\n\n端口未正确释放。")
        self.assertEqual(message["groups"][0]["thoughts"], ["我来分析这段代码。", "让我先梳理代码逻辑。"])

    async def test_thread_events_and_state_are_persistent(self):
        engine = create_engine("sqlite:///:memory:")
        Base.metadata.create_all(engine)
        session_factory = sessionmaker(bind=engine)
        conversation_id = str(uuid.uuid4())
        db = session_factory()
        user = UserModel(username="local-user")
        db.add(user)
        db.flush()
        db.add(ConversationModel(id=conversation_id, user_id=user.id, title="可恢复对话"))
        db.commit()
        db.close()

        with patch.object(main, "SessionLocal", session_factory):
            state = await main.set_conversation_state(conversation_id, main.ConversationStateRequest(is_pinned=True))
            events = await main.get_conversation_events(conversation_id)

        self.assertTrue(state["is_pinned"])
        self.assertEqual(events["events"][0]["type"], "thread_state_changed")
        self.assertEqual(events["events"][0]["sequence"], 1)

    async def test_conversation_can_be_renamed_and_deleted(self):
        engine = create_engine("sqlite:///:memory:")
        Base.metadata.create_all(engine)
        session_factory = sessionmaker(bind=engine)
        conversation_id = str(uuid.uuid4())
        db = session_factory()
        user = UserModel(username="local-user")
        db.add(user)
        db.flush()
        db.add(ConversationModel(id=conversation_id, user_id=user.id, title="旧任务名"))
        db.commit()
        db.close()

        with patch.object(main, "SessionLocal", session_factory):
            renamed = await main.rename_conversation(
                conversation_id,
                main.ConversationUpdateRequest(title="新任务名"),
            )
            deleted = await main.delete_conversation(conversation_id)

        self.assertEqual(renamed["title"], "新任务名")
        self.assertTrue(deleted["deleted"])
        db = session_factory()
        self.assertIsNone(db.get(ConversationModel, conversation_id))
        db.close()

    async def test_upload_keeps_original_attachment_and_timeline_event(self):
        import tempfile

        engine = create_engine("sqlite:///:memory:")
        Base.metadata.create_all(engine)
        session_factory = sessionmaker(bind=engine)
        conversation_id = str(uuid.uuid4())
        with tempfile.TemporaryDirectory() as temp_dir:
            upload = UploadFile(file=io.BytesIO("原始文件内容".encode()), filename="source.txt")
            with (
                patch.object(main, "SessionLocal", session_factory),
                patch.object(main, "ATTACHMENT_DIR", temp_dir),
            ):
                response = await main.upload_file(upload, conversation_id)
                payload = json.loads(response.body)
                events = await main.get_conversation_events(conversation_id)

            db = session_factory()
            attachment = db.get(AttachmentModel, payload["attachment_id"])
            self.assertIsNotNone(attachment)
            self.assertEqual(attachment.extracted_text, "原始文件内容")
            self.assertTrue(Path(attachment.storage_path).exists())
            self.assertEqual(Path(attachment.storage_path).read_text(encoding="utf-8"), "原始文件内容")
            self.assertEqual(events["events"][0]["type"], "attachment_added")
            db.close()


if __name__ == "__main__":
    unittest.main()
